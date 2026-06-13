from pathlib import Path
from shutil import copy2

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement


SOURCE = Path(
    r"C:\Users\alfie\OneDrive - JA GROUP SERVICES LTD\Documents"
    r"\JA Group Services Annual Business Review Report 2026 v2.docx"
)
OUTPUT = Path(
    r"C:\Users\alfie\OneDrive - JA GROUP SERVICES LTD\Documents"
    r"\JA Group Services Annual Business Review Report 2026 FINAL.docx"
)


REPLACEMENTS = {
    "Board Annual Business Review Report 2026": (
        "Annual Board Business Review Report 2026"
    ),
    "Jack and Alfie will review this report together, discuss the matters raised, "
    "agree any corrections or amendments required, and identify the next actions "
    "to be recorded in the Company governance records.": (
        "The current directors will review this report, discuss the matters raised, "
        "agree any corrections or amendments required, and identify the next actions "
        "to be recorded in the Company governance records. For this review, "
        "JA Group Services Ltd is represented by Mr Alfie Thomas Holywood Murray and "
        "by its corporate director, JSDS Group Ltd, acting through its director "
        "Mr Jack Nicolau Sousa Da Silva."
    ),
    "JA Group Services Ltd is approximately one year and three months into its "
    "corporate development as at June 2026.": (
        "JA Group Services Ltd (company number 16314179) was incorporated on "
        "13 March 2025 and is approximately one year and three months into its "
        "corporate development as at 8 June 2026."
    ),
    "The Board notes that the Company remains at an early stage of development and "
    "should continue to maintain clear records": (
        "The Company remains at an early stage of development and should continue "
        "to maintain clear records"
    ),
    "The Board has reviewed the Company current and developing website and service "
    "portfolio.": (
        "This report presents the Company’s current and developing website and "
        "service portfolio for Board review."
    ),
    "The Board considers JA Document Hub a significant direct service because it "
    "moves the Company beyond pure facilitation or referral activity.": (
        "JA Document Hub is proposed as a significant direct service because it "
        "moves the Company beyond pure facilitation or referral activity."
    ),
    "The service operates through the Company own website and portal environment, "
    "with customers engaging directly with JA Group Services Ltd through the "
    "platform.": (
        "The exported application contains a Company-operated website and portal "
        "environment intended for customers to engage directly with "
        "JA Group Services Ltd. Production deployment, live integrations and "
        "end-to-end customer journeys must be verified separately before launch."
    ),
    "Board feedback: JA Document Hub is strategically aligned with the Company "
    "updated 2026 direction and should remain a priority direct service, subject to "
    "continuing governance, data protection, payment and operational controls.": (
        "Proposed Board position: JA Document Hub is strategically aligned with the "
        "Company’s updated 2026 direction and should remain a priority direct "
        "service, subject to deployment testing and continuing governance, data "
        "protection, payment and operational controls."
    ),
    "The Board considers JA Smart Profile to be one of the Company priority direct "
    "services. It has its own pricing structure and customer account journey, "
    "supporting the Company transition to subscription-enabled digital service "
    "provision.": (
        "JA Smart Profile is proposed as one of the Company’s priority direct "
        "services. The exported application contains customer account, profile, "
        "pricing and billing functionality. Its legal wording currently describes "
        "paid plans as coming soon, so no paid plan should be described as live "
        "until production configuration and customer charging have been approved "
        "and tested."
    ),
    "Board feedback: JA Smart Profile should continue to be developed and operated, "
    "provided that public wording remains accurate, customer claims are not "
    "misleading, pricing is clear, customer support information is available and "
    "data protection requirements are maintained.": (
        "Proposed Board position: JA Smart Profile should continue to be developed, "
        "provided that public wording remains accurate, customer claims are not "
        "misleading, pricing is clear, customer support information is available "
        "and data protection requirements are maintained."
    ),
    "JA Booking is a direct booking portal and online appointment service developed "
    "by JA Group Services Ltd.": (
        "JA Booking is a developing direct booking portal and online appointment "
        "service of JA Group Services Ltd."
    ),
    "The Board notes that JA Booking pricing model is to be remodelled.": (
        "JA Booking’s pricing model is to be remodelled. The exported application "
        "marks paid plans as coming soon and contains demonstration data; these "
        "elements must not be treated as live customer, booking, payment or revenue "
        "records."
    ),
    "Board feedback: JA Booking should remain within the Company direct service "
    "development programme, but final pricing approval should be reserved until the "
    "pricing model has been remodelled and reviewed.": (
        "Proposed Board position: JA Booking should remain within the Company’s "
        "direct service development programme, but launch and final pricing approval "
        "should be reserved until the pricing model, production data, payment "
        "journey and operational controls have been remodelled, tested and reviewed."
    ),
    "The Board notes the development of JA Print Studio as a print-related service, "
    "brand or service project connected with JA Group Services Ltd.": (
        "JA Print Studio is a developing print-related service, brand or service "
        "project of JA Group Services Ltd."
    ),
    "Board feedback: JA Print Studio should be included in the annual business "
    "review as a developing service area, but it should receive a separate launch "
    "or operating review before full public commercial operation.": (
        "Proposed Board position: JA Print Studio should remain a developing service "
        "area and receive a separate launch and operating review before full public "
        "commercial operation."
    ),
    "JA Group Services Ltd has established JA Group Services Secure Access as the "
    "Company customer-facing authentication identity for its direct service "
    "websites.": (
        "JA Group Services Secure Access is the Company’s designated "
        "customer-facing authentication model for its direct service websites."
    ),
    "JA Group Services Secure Access is powered by Microsoft Entra External ID and "
    "provides secure customer authentication and a single sign-in identity for "
    "customers using direct Company services.": (
        "The exported applications contain Microsoft Entra External ID integration "
        "for customer authentication and a separate workforce-tenant route for "
        "administrative access. Live tenant configuration, redirect URIs, secrets, "
        "role assignments and end-to-end sign-in must be verified for each service "
        "before production use."
    ),
    "Customers will be prompted to create an account through JA Group Services "
    "Secure Access before accessing the Company direct service websites.": (
        "Where a direct service requires an account, customers should be directed "
        "through JA Group Services Secure Access once that service’s production "
        "authentication configuration has passed testing."
    ),
    "Stripe webhooks, APIs and related payment or billing settings have been "
    "configured by Mr Alfie Thomas Holywood Murray inside the relevant website admin "
    "portals. Stripe is intended to support payment, billing and subscription "
    "functions where applicable.": (
        "The exported applications contain Stripe-related APIs, webhook handlers, "
        "billing interfaces and administrative configuration areas. Mr Alfie Thomas "
        "Holywood Murray is responsible for the relevant Microsoft and payment "
        "administration. The source exports do not prove that production keys, "
        "webhook secrets, products, prices or live-mode settings are complete; these "
        "must be verified in Stripe and in each deployed service before any customer "
        "is charged."
    ),
    "All Company websites, direct service websites, customer portals and relevant "
    "public-facing web services must be checked before being published onto the "
    "mainstream web.": (
        "All Company websites, direct service websites, customer portals and relevant "
        "public-facing web services must undergo documented security review before "
        "launch and appropriate checks again after deployment."
    ),
    "Before publication, each website must be checked using the National Cyber "
    "Security Centre Check Cyber Security service: "
    "https://checkcybersecurity.service.ncsc.gov.uk/ip-check/form.": (
        "As part of launch assurance, each deployed public website or relevant "
        "public-facing IP address should be checked using the National Cyber Security "
        "Centre’s Check Your Cyber Security service: "
        "https://checkcybersecurity.service.ncsc.gov.uk/ip-check. Because this is a "
        "remote check of public-facing systems, it may need to be completed or "
        "repeated after deployment."
    ),
    "The purpose of this requirement is to support basic cybersecurity assurance "
    "before any website is made publicly available or treated as ready for "
    "mainstream public access.": (
        "This supports basic cybersecurity assurance but does not replace secure "
        "configuration, vulnerability management, authentication testing, payment "
        "testing, data protection review or professional security testing where the "
        "risk requires it."
    ),
    "JA Group Services Ltd is the data controller unless stated otherwise. The "
    "Company Data Protection Officer is Mr Alfie Thomas Holywood Murray. Data "
    "protection enquiries should be directed to "
    "dataprotection@jagroupservices.co.uk.": (
        "JA Group Services Ltd is the data controller unless stated otherwise. "
        "Mr Alfie Thomas Holywood Murray is the Company’s designated data protection "
        "contact, and enquiries should be directed to "
        "dataprotection@jagroupservices.co.uk. The Company should use the formal "
        "title “Data Protection Officer” only if it has documented that appointment, "
        "confirmed the role’s independence and absence of conflicts, and completed "
        "any required notification to the Information Commissioner’s Office."
    ),
    "The Board has considered the key risks arising from the Company current "
    "position and updated business direction.": (
        "This report identifies the key risks arising from the Company’s current "
        "position and updated business direction for Board consideration."
    ),
    "The following discussion points are included to support the joint Board review "
    "by Mr Alfie Thomas Holywood Murray and Mr Jack Nicolau Sousa Da Silva.": (
        "The following discussion points support review by Mr Alfie Thomas Holywood "
        "Murray and the corporate director, JSDS Group Ltd, represented for this "
        "purpose by its director Mr Jack Nicolau Sousa Da Silva."
    ),
    "The Board confirms that direct digital services are now the Company main "
    "business priority for 2026.": (
        "Subject to Board approval, direct digital services will remain the "
        "Company’s main business priority for 2026."
    ),
    "The Board notes that the Company is reviewing the future development of its "
    "tours.jagroupservices.co.uk page.": (
        "This report notes that the Company is reviewing the future development of "
        "its tours.jagroupservices.co.uk page."
    ),
    "The Board notes that JA Domain Hub remains an active service area of "
    "JA Group Services Ltd.": (
        "This report notes that JA Domain Hub remains an active service area of "
        "JA Group Services Ltd."
    ),
    "The Board considers this acceptable for the time being because JA Domain Hub "
    "operates through GoDaddy/SecureServer reseller infrastructure.": (
        "The proposed Board position is that this is acceptable for the time being "
        "because JA Domain Hub operates through GoDaddy/SecureServer reseller "
        "infrastructure."
    ),
    "The Board notes that the Company has developed materially during its early "
    "operating period.": (
        "This report records that the Company has developed materially during its "
        "early operating period."
    ),
    "Recently joined as an affiliate partner with a reported commission arrangement "
    "of 50%, subject to applicable affiliate terms and partner records.": (
        "Recently joined as an affiliate partner under a reported commission "
        "arrangement, subject to the applicable affiliate terms and retained partner "
        "records."
    ),
    "Company current direction": "Company’s current direction",
    "Company main development priority": "Company’s main development priority",
    "Company main strategic focus": "Company’s main strategic focus",
    "Company wider commercial model": "Company’s wider commercial model",
    "Company updated 2026 model": "Company’s updated 2026 model",
    "Company direct service websites": "Company’s direct service websites",
    "Company development": "Company’s development",
    "Company annual governance": "Company’s annual governance",
    "Board feedback:": "Proposed Board position:",
}


def iter_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    seen_cells = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested in cell.tables:
                    for nested_row in nested.rows:
                        for nested_cell in nested_row.cells:
                            for paragraph in nested_cell.paragraphs:
                                yield paragraph


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return False

    full_text = "".join(run.text for run in paragraph.runs)
    if old not in full_text:
        # Hyperlink text is included in paragraph.text but not paragraph.runs.
        # Rebuild only these exceptional paragraphs, retaining the paragraph style.
        paragraph.text = paragraph.text.replace(old, new)
        return True

    replaced = full_text.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced)
    return True


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    counts = {old: 0 for old in REPLACEMENTS}
    for paragraph in iter_paragraphs(document):
        for old, new in REPLACEMENTS.items():
            if replace_in_paragraph(paragraph, old, new):
                counts[old] += 1

    # Use the existing blank paragraph below the delivery metadata for status.
    status_paragraph = document.paragraphs[5]
    status_paragraph.text = "Document status: Draft for Board review and approval"
    status_paragraph.style = document.paragraphs[2].style
    for run in status_paragraph.runs:
        run.bold = True
        run.font.size = Pt(10)

    # The source used fixed manual page breaks between content groups. Once the
    # corrected wording expanded, those breaks produced empty pages. Let Word
    # paginate the tables naturally instead.
    for paragraph in document.paragraphs:
        for br in paragraph._p.xpath(".//w:br[@w:type='page']"):
            br.getparent().remove(br)

    # Keep the data-protection section and action-log table from producing
    # unlabelled continuation fragments on the following page.
    for table_index in (24, 30):
        table = document.tables[table_index]
        for row in table.rows:
            row_properties = row._tr.get_or_add_trPr()
            if not row_properties.xpath("./w:cantSplit"):
                row_properties.append(OxmlElement("w:cantSplit"))

    action_table = document.tables[30]
    action_paragraphs = [
        paragraph
        for row in action_table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    ]
    for paragraph in action_paragraphs[:-1]:
        paragraph.paragraph_format.keep_with_next = True

    risk_table = document.tables[26].cell(0, 0).tables[0]
    header_properties = risk_table.rows[0]._tr.get_or_add_trPr()
    if not header_properties.xpath("./w:tblHeader"):
        header_properties.append(OxmlElement("w:tblHeader"))
    for row in risk_table.rows:
        row_properties = row._tr.get_or_add_trPr()
        if not row_properties.xpath("./w:cantSplit"):
            row_properties.append(OxmlElement("w:cantSplit"))

    document.core_properties.title = "JA Group Services Ltd Annual Board Business Review Report 2026"
    document.core_properties.subject = "Draft for Board review and approval"
    document.core_properties.comments = (
        "Corrected against Companies House records and supplied application exports "
        "on 8 June 2026."
    )
    document.save(OUTPUT)

    missing = [old for old, count in counts.items() if count == 0]
    print(f"Saved: {OUTPUT}")
    print(f"Applied replacements: {sum(counts.values())}/{len(REPLACEMENTS)}")
    if missing:
        print("Missing replacements:")
        for item in missing:
            print(f"- {item}")
        raise SystemExit(1)


if __name__ == "__main__":
    main()
